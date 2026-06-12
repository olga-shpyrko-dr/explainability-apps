"""
Generates docs/IL_Explainability_App_User_Guide.docx
following DataRobot brand guidelines.
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ──────────────────────────────────────────────────────────────────
GREEN   = RGBColor(0x81, 0xFB, 0xA5)
BLACK   = RGBColor(0x0B, 0x0B, 0x0B)
INDIGO  = RGBColor(0x5C, 0x41, 0xFF)
GREY    = RGBColor(0xE4, 0xE4, 0xE4)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
DARK    = RGBColor(0x22, 0x22, 0x22)

ROOT = Path(__file__).parent.parent
OUT  = ROOT / "docs" / "IL_Explainability_App_User_Guide.docx"
OUT.parent.mkdir(exist_ok=True)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, name: str, size_pt: float, bold=False, color: RGBColor | None = None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def cell_bg(cell, color: RGBColor):
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_footer(doc: Document):
    section = doc.sections[0]
    footer  = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("© 2026 DataRobot, Inc. All rights reserved.   |   CONFIDENTIAL")
    set_font(run, "Courier New", 8, color=RGBColor(0x88, 0x88, 0x88))


def add_eyebrow(doc: Document, text: str):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    set_font(run, "Courier New", 8, color=RGBColor(0x66, 0x66, 0x66))
    return p


def add_title(doc: Document, text: str):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, "DM Sans", 20, bold=True, color=BLACK)
    return p


def add_subtitle(doc: Document, text: str):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    set_font(run, "DM Sans", 12, color=RGBColor(0x44, 0x44, 0x44))
    return p


def h1(doc: Document, text: str):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, "DM Sans", 16, bold=True, color=BLACK)
    # Green underline bar
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "81FBA5")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def h2(doc: Document, text: str):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, "DM Sans", 13, bold=True, color=BLACK)
    return p


def body(doc: Document, text: str):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, "DM Sans", 11, color=DARK)
    return p


def bullet(doc: Document, text: str, level=0):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, "DM Sans", 11, color=DARK)
    return p


def code_block(doc: Document, lines: list[str]):
    for line in lines:
        p   = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        run = p.add_run(line if line else " ")
        set_font(run, "Courier New", 9, color=RGBColor(0x1a, 0x1a, 0x1a))
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"),   "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"),  "F4F4F4")
        p._p.get_or_add_pPr().append(shading)


def table_row(tbl, cells: list[str], header=False):
    row = tbl.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        if header:
            set_font(run, "DM Sans", 10, bold=True, color=BLACK)
            cell_bg(cell, GREEN)
        else:
            set_font(run, "DM Sans", 10, color=DARK)
    return row


# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width   = Inches(8.5)
section.page_height  = Inches(11)
section.left_margin  = Inches(0.7)
section.right_margin = Inches(1.0)
section.top_margin   = Inches(0.7)
section.bottom_margin= Inches(1.0)

add_footer(doc)

# ── Cover ─────────────────────────────────────────────────────────────────────
add_eyebrow(doc, "JUNE 2026  ·  INTERNAL")
add_title(doc, "IL Protection Lapse Explainability App")
add_subtitle(doc, "User Guide — Python Backend & React Frontend")

body(doc,
     "This guide covers installation, configuration, and day-to-day use of the "
     "Explainability App for the Protection Lapse Propensity model. The app surfaces "
     "grouped SHAP explanations, cohort filtering, individual policy look-up, and an "
     "AI-generated narrative layer on top of DataRobot prediction outputs.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Architecture Overview")

body(doc,
     "The application follows a two-tier architecture: a Python FastAPI backend that "
     "fetches live data from DataRobot at startup, and a lightweight React + Vite "
     "frontend that consumes its REST API.")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
table_row(tbl, ["Layer", "Technology", "Purpose"], header=True)
for row in [
    ("Backend", "Python 3.11 · FastAPI · DataRobot SDK", "Data pipeline, SHAP aggregation, LLM narrative, REST API"),
    ("Frontend", "React 18 · Vite · Recharts · TypeScript", "Cohort filter sidebar, group chart, policy look-up, narrative UI"),
    ("DataRobot", "DataRobot EU (app.eu.datarobot.com)", "Prediction Explanations source, LLM Gateway"),
    ("SQL export", "SQLAlchemy · pyodbc · SQL Server", "Power BI flat-table export (optional)"),
]:
    table_row(tbl, list(row))

doc.add_paragraph()
body(doc,
     "On startup the backend uploads the scoring dataset to the project, requests "
     "predictions, initialises Prediction Explanations, and loads the results into "
     "memory. A local cache (.prediction_dataset_cache.json) skips the upload step "
     "on subsequent restarts.")

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Prerequisites")

h2(doc, "2.1 Python environment")
for item in [
    "Python 3.11 or later",
    "pip install -r requirements.txt   (from project root)",
    "ODBC Driver 17 for SQL Server — only required for the SQL export script",
]:
    bullet(doc, item)

h2(doc, "2.2 DataRobot access")
for item in [
    "DataRobot EU account with access to project 6a22f2218f74af009899ddb1",
    "API token with read access to Projects, Datasets, Deployments, and LLM Gateway",
    "The scoring dataset (ID 6a2275eb326d5530a77a0b30) must be in the AI Catalog",
]:
    bullet(doc, item)

h2(doc, "2.3 Node.js (frontend only)")
bullet(doc, "Node.js 18 or later and npm")

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "3. Configuration")

body(doc,
     "Copy .env.example to .env in the project root and fill in the required values. "
     "The backend reads this file via pydantic-settings; all keys are case-insensitive.")

h2(doc, "3.1 Required keys")
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
table_row(tbl, ["Key", "Example value", "Description"], header=True)
for row in [
    ("DATAROBOT_API_TOKEN", "abc123…", "DataRobot personal API token"),
    ("DATAROBOT_ENDPOINT", "https://app.eu.datarobot.com/api/v2", "DataRobot EU REST endpoint"),
]:
    table_row(tbl, list(row))

doc.add_paragraph()
h2(doc, "3.2 Optional keys")
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
table_row(tbl, ["Key", "Example value", "Description"], header=True)
for row in [
    ("DR_GATEWAY_MODEL", "azure/gpt-4o-mini", "LLM Gateway model ID (see Section 5)"),
    ("DR_LLM_DEPLOYMENT_ID", "<deployment id>", "DataRobot deployed TextGen model"),
    ("AZURE_OPENAI_API_KEY", "<key>", "Azure OpenAI key"),
    ("AZURE_OPENAI_API_BASE", "https://…openai.azure.com/", "Azure OpenAI endpoint"),
    ("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o", "Azure deployment name"),
    ("ANTHROPIC_API_KEY", "<key>", "Anthropic direct API key"),
    ("SQL_CONNECTION_STRING", "mssql+pyodbc://…", "SQL Server connection (export only)"),
    ("TRAINING_EXCEL_PATH", "data/Training.xlsx", "Adds Lapse_ind labels to export"),
]:
    table_row(tbl, list(row))

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "4. Starting the Application")

h2(doc, "4.1 Backend")
body(doc, "From the project root:")
code_block(doc, [
    "cd backend",
    "source ../.venv/bin/activate        # or .venv\\Scripts\\activate on Windows",
    "uvicorn main:app --reload --port 8000",
])
body(doc,
     "First startup takes 2–5 minutes while the prediction dataset is uploaded and "
     "Prediction Explanations are computed. Progress is logged to the terminal. "
     "Subsequent starts are fast — the dataset ID is cached in "
     ".prediction_dataset_cache.json.")

h2(doc, "4.2 Frontend")
code_block(doc, [
    "cd frontend",
    "npm install       # first time only",
    "npm run dev",
])
body(doc, "Open http://localhost:5173 in your browser.")

h2(doc, "4.3 Health check")
body(doc, "Verify the backend is ready:")
code_block(doc, ["curl http://localhost:8000/api/health"])
body(doc,
     "A successful response includes rows_loaded (should be ~6,668), explanation_rows "
     "(~26,672), and sample row IDs.")

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "5. LLM Narrative Configuration")

body(doc,
     "The AI Narrative tab supports four LLM providers. Configure at least one to "
     "enable narrative generation. The app automatically detects which providers have "
     "credentials and shows only those as selectable options.")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
table_row(tbl, ["Provider", "Required .env keys", "Notes"], header=True)
for row in [
    ("DR LLM Gateway", "DR_GATEWAY_MODEL", "Recommended. Uses your DataRobot API token — no extra credentials needed."),
    ("DR Deployed LLM", "DR_LLM_DEPLOYMENT_ID", "Calls a deployed TextGen model via its chat completions endpoint."),
    ("Azure OpenAI", "AZURE_OPENAI_API_KEY, AZURE_OPENAI_API_BASE, AZURE_OPENAI_DEPLOYMENT_NAME", "Standard Azure OpenAI credentials."),
    ("Anthropic", "ANTHROPIC_API_KEY", "Direct Anthropic API. Model defaults to claude-sonnet-4-6."),
]:
    table_row(tbl, list(row))

doc.add_paragraph()
h2(doc, "5.1 LLM Gateway model IDs")
body(doc,
     "Set DR_GATEWAY_MODEL to the Chat model ID from the DataRobot documentation "
     "(docs.datarobot.com > Reference > LLM Availability). Examples:")
for item in [
    "azure/gpt-4o-mini",
    "azure/gpt-4o-2024-11-20",
    "anthropic/claude-sonnet-4-6",
    "vertex_ai/gemini-2.0-flash-001",
    "bedrock/anthropic.claude-3-5-sonnet-20241022-v2:0",
]:
    bullet(doc, item)

body(doc,
     "After changing .env, restart the backend — settings are cached at startup "
     "and a hot-reload is not sufficient.")

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "6. API Endpoints")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
table_row(tbl, ["Method + Path", "Parameters", "Returns"], header=True)
for row in [
    ("GET /api/health", "—", "Status, row counts, sample row IDs"),
    ("GET /api/columns", "—", "All filterable columns with type, min/max, or distinct values"),
    ("GET /api/cohort", "filters (JSON string)", "Cohort size, score distribution, histogram"),
    ("GET /api/groups", "filters (JSON string)", "Group SHAP aggregations sorted by |avg_shap|"),
    ("GET /api/row/{row_id}", "—", "Waterfall explanation for one policy"),
    ("GET /api/llm/providers", "—", "Configured LLM providers and which are available"),
    ("POST /api/narrative", "JSON body: filters, provider, custom_instruction", "AI-generated narrative text"),
]:
    table_row(tbl, list(row))

doc.add_paragraph()
h2(doc, "6.1 Filter format")
body(doc, "Filters are passed as a JSON string in the query parameter. Examples:")
code_block(doc, [
    '# Range filter',
    '{"Age_life1": {"min": 30, "max": 50}}',
    '',
    '# Multi-select filter',
    '{"Product_Desc": ["Block Life Term Cover", "Whole of Life"]}',
    '',
    '# Combined',
    '{"Decile": {"min": 8, "max": 10}, "SmokerStatus": ["Y"]}',
])

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "7. Feature Group Configuration")

body(doc,
     "Feature-to-group assignments are stored in backend/feature_group_mapping.json. "
     "Edit this file to reassign features without touching any code. "
     "The backend reloads the mapping on restart.")

h2(doc, "7.1 Current groups")
for grp in [
    "Policy & Product — product type, cover, premium, term, sum assured",
    "Policy Portfolio — number and value of policies across lives",
    "Agent / Adviser — servicing agent, source of business, commission",
    "Sociodemographic — age, gender, smoker status, occupation, family",
    "Financial Profile — income, debt, net worth, credit score, monthly expenditure",
    "Engagement & Reviews — logins, review type, months since review",
    "Persona — persona segment labels",
]:
    bullet(doc, grp)

h2(doc, "7.2 Adding or moving a feature")
body(doc, "Open backend/feature_group_mapping.json and move the feature name string "
     "to the target group array. Restart the backend.")

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "8. SQL Server Export (Power BI Path)")

body(doc,
     "The script scripts/export_to_sql.py pulls the same data as the backend and "
     "writes three tables to SQL Server for use in Power BI.")

h2(doc, "8.1 Setup")
body(doc, "Add SQL_CONNECTION_STRING to .env:")
code_block(doc, [
    "# SQL auth",
    "SQL_CONNECTION_STRING=mssql+pyodbc://user:pass@server/db?driver=ODBC+Driver+17+for+SQL+Server",
    "",
    "# Windows auth",
    "SQL_CONNECTION_STRING=mssql+pyodbc://@server/db?driver=ODBC+Driver+17+for+SQL+Server&trusted_connection=yes",
])

h2(doc, "8.2 Running the export")
code_block(doc, [
    "# First run — create tables and populate",
    "python scripts/export_to_sql.py --replace",
    "",
    "# Include training labels (Lapse_ind)",
    "python scripts/export_to_sql.py --replace --training-excel data/Training.xlsx",
])

h2(doc, "8.3 Output tables")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
table_row(tbl, ["Table", "Contents"], header=True)
for row in [
    ("explainability_scored_population", "All source features + prediction score (~6,668 rows)"),
    ("explainability_explanation_long", "Unpivoted SHAP rows with group labels (~26,672 rows)"),
    ("explainability_feature_group_mapping", "Feature → group lookup (83 rows)"),
]:
    table_row(tbl, list(row))

doc.add_paragraph()
h2(doc, "8.4 Power BI setup")
for item in [
    "Connect to SQL Server (Import or DirectQuery)",
    "Load all three tables",
    "Create relationship: scored_population[Policy_Number]  →  explanation_long[row_id]  (1 : many)",
    "Add DAX measure:  Avg Group SHAP = AVERAGEX(RELATEDTABLE(explanation_long), explanation_long[shap_strength])",
    "Add slicers on: Decile, Product_Desc, Age_life1, SmokerStatus, feature_group",
]:
    bullet(doc, item)

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "9. Troubleshooting")

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
table_row(tbl, ["Symptom", "Fix"], header=True)
for row in [
    ("Startup takes > 10 min", "Check that .prediction_dataset_cache.json exists. If missing, the full PE pipeline reruns. Delete the cache only if the scoring dataset has changed."),
    ("No explanation data for cohort", "Run curl /api/health and check explanation_sample_row_ids matches population_sample_row_ids format. A type mismatch causes the join to fail."),
    ("LLM narrative returns 500", "Check backend logs for the LiteLLM error. Most common causes: wrong DR_GATEWAY_MODEL format (must be azure/gpt-4o-mini not azure-openai/gpt-4o-mini), or backend not restarted after .env change."),
    ("Frontend shows stale data", "The backend caches data at startup. Restart uvicorn to pull fresh explanations from DataRobot."),
    ("SQL export fails with driver error", "Install ODBC Driver 17 for SQL Server. On macOS: brew install msodbcsql17."),
]:
    table_row(tbl, list(row))

# ═════════════════════════════════════════════════════════════════════════════
h1(doc, "10. Key File Reference")

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
table_row(tbl, ["File / Directory", "Purpose"], header=True)
for row in [
    (".env", "Runtime secrets and configuration (not committed to git)"),
    ("backend/main.py", "FastAPI app, lifespan startup, all REST endpoints"),
    ("backend/pipeline.py", "DataRobot data pipeline: upload → predict → PE → load"),
    ("backend/cohort.py", "Filter engine, cohort profile stats, group SHAP aggregation"),
    ("backend/llm_client.py", "LiteLLM abstraction for all four LLM providers"),
    ("backend/narrative.py", "Prompt construction and LLM call for narrative generation"),
    ("backend/config.py", "pydantic-settings config model — all environment variables"),
    ("backend/feature_group_mapping.json", "Feature → group assignments (editable without code changes)"),
    ("backend/.prediction_dataset_cache.json", "Auto-generated PE dataset ID cache — delete to force re-upload"),
    ("frontend/src/components/", "React UI components: CohortFilter, GroupExplanationChart, WaterfallChart, NarrativePanel"),
    ("scripts/export_to_sql.py", "SQL Server export for Power BI consumption"),
    ("docs/", "This document and the original app specification"),
]:
    table_row(tbl, list(row))

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved → {OUT}")
